from docx import Document


documento = Document()
#Titulo
documento.add_heading('Relatório Mensal de Vendas',0)
#Paragrafo
paragrafo = documento.add_paragraph('Neste mês de agosto foram realizadas um total de 10 vendas de notebooks. Segue em anexo a tabela com os dados de cada venda realizada.')
#Titulo tabela
subtitle = documento.add_paragraph('Vendas de Agosto', style='Heading 1')
#Dados
dados =[
    ['Modelo 1','R$1500','28/08/2022','Robert','Loja 1'],
    ['Modelo 2','R$2500','28/08/2022','Amanda','Loja 1'],
    ['Modelo 2','R$2500','28/08/2022','Robert','Loja 1'],
    ['Modelo 3','R$3500','28/08/2022','Robert','Loja 1'],
    ['Modelo 3','R$3500','28/08/2022','Carol','Loja 1'],
    ['Modelo 3','R$3500','28/08/2022','Amanda','Loja 1'],
    ['Modelo 4','R$5000','28/08/2022','Robert','Loja 1'],
    ['Modelo 4','R$5000','28/08/2022','Amanda','Loja 1'],
    ['Modelo 4','R$5000','28/08/2022','Robert','Loja 1'],
    ['Modelo 4','R$5000','28/08/2022','Amanda','Loja 1']
]

#Cabeçalho da Tabela
tabela = documento.add_table(rows=1,cols=5)
cabecalho = tabela.rows[0].cells
cabecalho[0].text = 'Modelo'
cabecalho[1].text = 'Preço'
cabecalho[2].text = 'Data'
cabecalho[3].text = 'Vendedor'
cabecalho[4].text = 'Loja'

# Dados da Tabela
for modelo, preco, data, vendedor, loja in dados:
    linha_atual = tabela.add_row().cells
    linha_atual[0].text = modelo
    linha_atual[1].text = preco
    linha_atual[2].text = data
    linha_atual[3].text = vendedor
    linha_atual[4].text = loja

paragrafo = documento.add_paragraph('')
paragrafo = documento.add_paragraph('Para as vendas deste mês, o funcionário Robert foi o funcionário com a maior quantidade de vendas diretas.')
paragrafo = documento.add_paragraph('O lucro total gerado na loja 1 foi de R$37000.00')

#Salvar arquivo 
documento.save('Projeto01.docx')