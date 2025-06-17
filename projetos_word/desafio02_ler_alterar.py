''' DESAFIO 
* Altere o texto:
He Had Decided On A New Policy.
Para:
esse texto foi alterado

* Altere o texto:
A tagline for an airline: Take the High Road. Rumour has it targeted online advertising was developed because the internet was upset that you could read it but it couldn't read you. Trepidelicious. I'm in a band that does Metallica covers with our private parts - it's called Myphallica. Petrovache. Pantone is a colour but also the singular version of pants. Most streets are two-way streets...why does that make love so special?.
Para:
Este parágrafo foi alterado

* Altere o texto:
Nice to meet you too
Para:
Good morning'''

from docx import Document

documento = Document('sampleEnchanced.docx')

# Achando o "índice" e o texto do documento
for index, texto in enumerate(documento.paragraphs):
    print('-'*30)
    print(index)
    print(texto.text)

# Alterações
documento.paragraphs[10].text = 'esse texto foi alterado '

documento.paragraphs[22].text = 'Esse parágrafo foi alterado '

documento.paragraphs[33].text = 'Good morning'

# Salvando alterações em um novo arquivo
documento.save('desafio_ler_alterar.docx')
